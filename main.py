import os
import json
from pathlib import Path
import re
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from openai import OpenAI
from pydantic import BaseModel
from docx import Document

from validation import ValidationError, parse_and_validate_project_json

app = FastAPI()
load_dotenv()
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "contract_template_demo.docx"
OUTPUT_PATH = BASE_DIR / "contractfinal.docx"
PROJECT_JSON_PATH = BASE_DIR / "project_characteristics.json"
UNUSED_FIELDS_PATH = BASE_DIR / "contract_unused_fields.json"

print(os.getenv("OPENAI_API_KEY"))
OPENAI_MODEL = "gpt-5-mini"
OPENAI_MAX_RETRIES = 3
PROJECT_FACTS = """
Reference contract constraints (must stay within these ranges):

- Product type: software or mobile application development.
- Work stages count: from 3 to 4 stages.
- Total duration: typically from 60 to 180 working days.
- Prepayment: from 30% to 50% of first stage payment.
- Stage payment percentages: must sum to 100%.
- Intellectual property model: exclusive rights transfer to the client.
- Source code transfer: required after full payment.
- Access credentials/accounts: transferred to the client upon project completion.
- Delay penalty per day: from 0.1% to 0.5%.
- Penalty cap: exactly 10%.
- Warranty period: from 2 to 6 months.
"""

SYSTEM_PROMPT = """You are a system analyst in software development.

Your task is to analyze a brief project description and return strictly valid JSON with project parameters for generating a software development contract.

Do not generate a contract.
Do not provide explanations.
The response must contain only JSON, with no text outside the JSON.

If data is insufficient, make professional assumptions based on standard software development practices in Russia,
while strictly staying within the allowed contract constraint ranges provided in the reference characteristics.

Forbidden:

- adding any text outside JSON
- modifying the JSON structure
- adding or removing fields
- violating required schema constraints
- generating values outside the specified numeric ranges

Numeric requirements:

- total_duration_working_days — integer, must be between 60 and 180 (days)
- stages_count — integer, must be 3 or 4
- number of objects in "stages" array must equal stages_count
- stage payment percentages — must sum exactly to 100
- stage duration sum — must equal total_duration_working_days
- prepayment_percent — integer, must be between 30 and 50
- penalty_percent_per_day — number, must be between 0.1 and 0.5
- penalty_cap_percent — exactly 10
- warranty_claim_window_months — integer, must be between 2 and 6

Contract logic requirements:

- ip_transfer_model must always be "exclusive_transfer"
- access_transfer_required must always be true
- product_type must correspond to software or mobile application development

Required JSON structure:
{
  "project_summary": "",
  "product_type": "",
  "work_scope_items": [""],
  "total_duration_working_days": 0,
  "stages_count": 0,
  "stages": [
    {
      "name": "",
      "duration_working_days": 0,
      "payment_percent": 0
    }
  ],
  "prepayment_percent": 0,
  "ip_transfer_model": "exclusive_transfer",
  "access_transfer_required": true,
  "penalty_percent_per_day": 0.0,
  "penalty_cap_percent": 10,
  "warranty_claim_window_months": 0
}

Use explicit data from the description if provided.
If explicit values contradict allowed ranges, adjust them to the nearest valid value within constraints.
Otherwise, estimate independently but remain realistic and internally consistent.
"""


class ContractData(BaseModel):
    contract_number: str
    city: str
    contract_day: str
    contract_month: str
    contract_year: str

    customer_company_name: str
    customer_representative_name: str
    customer_representative_basis: str
    customer_inn: str
    customer_ogrn_or_ogrnip: str
    customer_legal_address: str
    customer_bank: str
    customer_bik: str
    customer_correspondent_account: str
    customer_settlement_account: str
    customer_kpp: str = "Не указано"

    contractor_type: str
    contractor_company_name: str
    contractor_representative_name: str
    contractor_representative_basis: str
    contractor_inn: str
    contractor_ogrn_or_ogrnip: str
    contractor_legal_address: str
    contractor_bank: str
    contractor_bik: str
    contractor_correspondent_account: str
    contractor_settlement_account: str

    vat_type: str
    price_value: str = "Не указана"
    project_description: str = ""


# -----------------------
# TEMPLATE ENGINE
# -----------------------

def extract_template_variables(template_text: str) -> set[str]:
    return set(re.findall(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}", template_text))


def inflect_word(word: str, case: str) -> str:
    """
    Простое склонение русских ФИО без библиотек.
    case: gent (Р.п.), datv (Д.п.), ablt (Т.п.)
    """

    if not word:
        return word

    # -------- ФАМИЛИИ --------
    if word.endswith("ов") or word.endswith("ев") or word.endswith("ин"):
        if case == "gent":
            return word + "а"
        if case == "datv":
            return word + "у"
        if case == "ablt":
            return word + "ым"

    if word.endswith("ова") or word.endswith("ева") or word.endswith("ина"):
        if case == "gent":
            return word[:-1] + "ой"
        if case == "datv":
            return word[:-1] + "ой"
        if case == "ablt":
            return word[:-1] + "ой"

    # -------- ИМЕНА --------
    if word.endswith("й"):
        base = word[:-1]
        if case == "gent":
            return base + "я"
        if case == "datv":
            return base + "ю"
        if case == "ablt":
            return base + "ем"

    if word.endswith("а"):
        base = word[:-1]
        if case == "gent":
            return base + "ы"
        if case == "datv":
            return base + "е"
        if case == "ablt":
            return base + "ой"

    if word.endswith("я"):
        base = word[:-1]
        if case == "gent":
            return base + "и"
        if case == "datv":
            return base + "е"
        if case == "ablt":
            return base + "ей"

    # -------- ОТЧЕСТВА --------
    if word.endswith("ич"):
        if case == "gent":
            return word + "а"
        if case == "datv":
            return word + "у"
        if case == "ablt":
            return word + "ем"

    if word.endswith("на"):
        base = word[:-1]
        if case == "gent":
            return base + "ы"
        if case == "datv":
            return base + "е"
        if case == "ablt":
            return base + "ой"

    # fallback
    return word


def inflect_fio_case(full_name: str, case: str) -> str:
    if not full_name:
        return ""

    if case == "nomn":
        return full_name

    parts = full_name.split()
    inflected = [inflect_word(p, case) for p in parts]
    return " ".join(inflected)


def to_initials(full_name: str) -> str:
    parts = [part for part in full_name.split() if part]
    if not parts:
        return ""
    surname = parts[0]
    initials = "".join(f"{p[0]}." for p in parts[1:] if p)
    return f"{surname} {initials}".strip()


def build_context(payload: dict) -> dict:
    fio = payload.get("contractor_representative_name", "")
    customer_fio = payload.get("customer_representative_name", "")
    contractor_intro = f"{payload.get('contractor_type', '')} {fio}".strip()
    if payload.get("contractor_type") == "ИП":
        contractor_intro = f"ИП {to_initials(fio)}".strip()

    return {
        **payload,
        "product_genitive": "программного обеспечения",
        "customer_representative_name_genitive": inflect_fio_case(customer_fio, "gent"),
        "contractor_fio_full": fio,
        "contractor_ogrnip": payload.get("contractor_ogrn_or_ogrnip", ""),
        "customer_ogrn": payload.get("customer_ogrn_or_ogrnip", ""),
        "customer_kpp": payload.get("customer_kpp", "Не указано"),
        "contractor_address": payload.get("contractor_legal_address", ""),
        "contractor_fio_nominative": fio,
        "contractor_fio_genitive": inflect_fio_case(fio, "gent"),
        "contractor_fio_dative": inflect_fio_case(fio, "datv"),
        "contractor_fio_instrumental": inflect_fio_case(fio, "ablt"),
        "contractor_intro_name": contractor_intro,
    }


def enrich_context_with_project_characteristics(context: dict, project_characteristics: dict) -> dict:
    """Merge validated LLM JSON into template context and prepare rendered fields."""
    merged = {**context, **project_characteristics}

    # Human-friendly stage lines for template loop.
    stage_lines = []
    for idx, stage in enumerate(project_characteristics.get("stages", []), start=1):
        stage_lines.append(
            f"Этап {idx}: {stage['name']} — {stage['duration_working_days']} рабочих дней, оплата {stage['payment_percent']}%."
        )
    merged["stage_lines"] = stage_lines

    work_scope_items = project_characteristics.get("work_scope_items")
    if not isinstance(work_scope_items, list) or not work_scope_items:
        work_scope_items = ["Объем работ уточняется в техническом задании."]
    merged["work_scope_items"] = [str(item) for item in work_scope_items]

    product_type = str(project_characteristics.get("product_type", "")).lower()
    if "мобиль" in product_type:
        merged["product_genitive"] = "мобильного приложения"
    else:
        merged["product_genitive"] = "программного обеспечения"

    if project_characteristics.get("ip_transfer_model") == "exclusive_transfer":
        merged["ip_transfer_clause"] = (
            "Исключительные права на результат работ переходят к Заказчику после полной оплаты."
        )
    else:
        merged["ip_transfer_clause"] = "Порядок перехода прав определяется дополнительным соглашением Сторон."

    merged["access_transfer_clause"] = (
        "Доступы и учетные данные передаются Заказчику по завершении проекта."
        if project_characteristics.get("access_transfer_required")
        else "Передача доступов отдельно не требуется."
    )

    return merged


def extract_loop_variables(template_text: str) -> set[str]:
    return set(
        re.findall(
            r"\{%\s*for\s+([a-zA-Z_][a-zA-Z0-9_]*)\s+in\s+[a-zA-Z_][a-zA-Z0-9_]*\s*%\}",
            template_text,
        )
    )


def validate_template_coverage(template_vars: set[str], context: dict, loop_vars: set[str]) -> None:
    """Ensure all template placeholders are provided by merged context."""
    missing = sorted(name for name in template_vars if name not in context and name not in loop_vars)
    if missing:
        raise HTTPException(
            status_code=500,
            detail=(
                "Template placeholders do not match available payload/LLM fields. "
                f"Missing variables: {missing}"
            ),
        )


LEGACY_LLM_TEMPLATE_ALIASES = {
    "work_scope": "work_scope_items",
    "subject_clause": "project_summary",
}


def validate_llm_template_alignment(template_text: str) -> None:
    """Reject legacy placeholders so template keeps matching actual LLM JSON field names."""
    placeholders = extract_template_variables(template_text)
    loop_iterables = set(
        re.findall(
            r"\{%\s*for\s+[a-zA-Z_][a-zA-Z0-9_]*\s+in\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*%\}",
            template_text,
        )
    )

    referenced_names = placeholders.union(loop_iterables)
    alias_errors = {
        old: new
        for old, new in LEGACY_LLM_TEMPLATE_ALIASES.items()
        if old in referenced_names
    }
    if alias_errors:
        raise HTTPException(
            status_code=500,
            detail=(
                "Template uses outdated LLM placeholder names: "
                + ", ".join(f"{old} -> {new}" for old, new in sorted(alias_errors.items()))
            ),
        )


def iter_docx_text_containers(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def extract_docx_template_variables(document: Document) -> set[str]:
    vars_found: set[str] = set()
    for container in iter_docx_text_containers(document):
        vars_found.update(extract_template_variables(container.text))
    return vars_found


def render_docx_template(document: Document, context: dict) -> None:
    _render_paragraph_loop_blocks(document.paragraphs, context)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                _render_paragraph_loop_blocks(cell.paragraphs, context)


def _render_paragraph_loop_blocks(paragraphs, context: dict) -> None:
    start_re = re.compile(r"\{%\s*for\s+([a-zA-Z_][a-zA-Z0-9_]*)\s+in\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*%\}")
    end_re = re.compile(r"\{%\s*endfor\s*%\}")

    i = 0
    while i < len(paragraphs):
        text = paragraphs[i].text
        if "{%" not in text:
            paragraphs[i].text = render_template(text, context)
            i += 1
            continue

        start_match = start_re.search(text)
        if not start_match:
            paragraphs[i].text = render_template(text, context)
            i += 1
            continue

        loop_var = start_match.group(1)
        iterable_name = start_match.group(2)
        iterable = context.get(iterable_name, [])
        if not isinstance(iterable, list):
            iterable = []

        end_index = i
        while end_index < len(paragraphs) and not end_re.search(paragraphs[end_index].text):
            end_index += 1

        if end_index >= len(paragraphs):
            paragraphs[i].text = render_template(text, context)
            i += 1
            continue

        block_lines = [paragraphs[idx].text for idx in range(i, end_index + 1)]
        block_lines[0] = start_re.sub("", block_lines[0], count=1)
        block_lines[-1] = end_re.sub("", block_lines[-1], count=1)
        block_template = "\n".join(block_lines).strip("\n")

        rendered_chunks = []
        for value in iterable:
            rendered_chunks.append(render_template(block_template, {**context, loop_var: value}))
        paragraphs[i].text = "\n".join(chunk for chunk in rendered_chunks if chunk)

        for clear_idx in range(i + 1, end_index + 1):
            paragraphs[clear_idx].text = ""

        i = end_index + 1


def render_template(template_text: str, context: dict) -> str:
    rendered = template_text

    # {% for item in work_scope %}...{{ item }}...{% endfor %}
    def for_replacer(match: re.Match) -> str:
        loop_var = match.group(1)
        iterable_name = match.group(2)
        block = match.group(3)
        iterable = context.get(iterable_name, [])
        if not isinstance(iterable, list):
            return ""

        chunks = []
        for value in iterable:
            chunks.append(re.sub(r"\{\{\s*" + re.escape(loop_var) + r"\s*\}\}", str(value), block))
        return "".join(chunks)

    rendered = re.sub(
        r"\{%\s*for\s+([a-zA-Z_][a-zA-Z0-9_]*)\s+in\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*%\}(.*?)\{%\s*endfor\s*%\}",
        for_replacer,
        rendered,
        flags=re.DOTALL,
    )

    # {% if var == "value" %}
    def if_replacer(match: re.Match) -> str:
        var_name = match.group(1)
        expected_value = match.group(2)
        block = match.group(3)

        return block if str(context.get(var_name, "")) == expected_value else ""

    rendered = re.sub(
        r"\{%\s*if\s+([a-zA-Z_][a-zA-Z0-9_]*)\s*==\s*\"([^\"]+)\"\s*%\}(.*?)\{%\s*endif\s*%\}",
        if_replacer,
        rendered,
        flags=re.DOTALL,
    )

    # {{ variable }}
    def var_replacer(match: re.Match) -> str:
        name = match.group(1)
        return str(context.get(name, ""))

    rendered = re.sub(
        r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}",
        var_replacer,
        rendered,
    )

    return rendered


def _extract_response_text(response) -> str:
    if getattr(response, "output_text", None):
        return response.output_text
    return "".join(
        chunk.text
        for item in getattr(response, "output", [])
        if getattr(item, "type", "") == "message"
        for chunk in getattr(item, "content", [])
        if getattr(chunk, "type", "") == "output_text"
    )


def generate_project_characteristics(project_description: str) -> dict:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="OPENAI_API_KEY is not configured")

    client = OpenAI(api_key=api_key)
    total_tokens_spent = 0
    last_error = "Unknown error"

    for attempt in range(1, OPENAI_MAX_RETRIES + 1):
        response = client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {
                    "role": "user",
                    "content": (
                        "Project description:\n"
                        f"{project_description}\n\n"
                        "Contract reference characteristics:\n"
                        f"{PROJECT_FACTS}"
                    ),
                },
            ],
        )

        usage = getattr(response, "usage", None)
        if usage and getattr(usage, "total_tokens", None):
            total_tokens_spent += usage.total_tokens

        response_text = _extract_response_text(response)

        try:
            validated_payload = parse_and_validate_project_json(response_text)
            print(f"[OpenAI] total tokens spent: {total_tokens_spent}")
            return validated_payload
        except ValidationError as exc:
            last_error = str(exc)
            print(f"[OpenAI] validation failed on attempt {attempt}: {last_error}")

    print(f"[OpenAI] total tokens spent before failure: {total_tokens_spent}")
    raise HTTPException(status_code=422, detail=f"LLM JSON validation failed: {last_error}")


# -----------------------
# API
# -----------------------

@app.post("/generate-contract")
def generate_contract(data: ContractData):
    payload = data.model_dump()

    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="Template not found")

    try:
        document = Document(TEMPLATE_PATH)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    template_text = "\n".join(container.text for container in iter_docx_text_containers(document))
    template_vars = extract_docx_template_variables(document)
    loop_vars = extract_loop_variables(template_text)
    validate_llm_template_alignment(template_text)
    
    project_characteristics = generate_project_characteristics(payload.get("project_description", ""))
    context = enrich_context_with_project_characteristics(
        build_context(payload),
        project_characteristics,
    )

    validate_template_coverage(template_vars, context, loop_vars)
    render_docx_template(document, context)

    unused_context_keys = sorted(key for key in context.keys() if key not in template_vars)
    unused_context_data = {key: context[key] for key in unused_context_keys}

    try:
        document.save(OUTPUT_PATH)
        PROJECT_JSON_PATH.write_text(
            json.dumps(project_characteristics, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        UNUSED_FIELDS_PATH.write_text(
            json.dumps(unused_context_data, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    except OSError as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    warning_message = (
        "Часть переменных не использована в DOCX шаблоне. "
        "Они сохранены в отдельный JSON файл."
        if unused_context_keys
        else ""
    )

    return {
        "status": "ok",
        "output_file": OUTPUT_PATH.name,
        "project_json_file": PROJECT_JSON_PATH.name,
        "unused_fields_file": UNUSED_FIELDS_PATH.name,
        "unused_context_keys": unused_context_keys,
        "unused_context_warning": bool(unused_context_keys),
        "warning_message": warning_message,
        "morphology": "disabled",
        "project_characteristics": project_characteristics,
    }
